"""One-off converter: a markdown doc -> a .docx with matching structure.

Usage: python scripts/md_to_docx.py <src.md> <dst.docx>

Not a general Markdown renderer; handles exactly the subset docs/ uses
(headings, bullets, tables, bold/italic/code/links). Re-run after editing
the source .md so the .docx stays in sync with it.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|\[.+?\]\(.+?\))")


def add_inline_runs(paragraph, text: str) -> None:
    for chunk in INLINE_RE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            paragraph.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            paragraph.add_run(chunk[1:-1])
        elif chunk.startswith("[") and "](" in chunk and chunk.endswith(")"):
            label, url = chunk[1:-1].split("](", 1)
            run = paragraph.add_run(f"{label} ({url})")
            run.italic = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            paragraph.add_run(chunk[1:-1]).italic = True
        else:
            paragraph.add_run(chunk)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        if not re.match(r"^\|[\s:|\-]+\|$", lines[i].strip()):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            rows.append(cells)
        i += 1
    return rows, i


def build(doc: Document, lines: list[str]) -> None:
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip("\n")

        if not stripped.strip():
            i += 1
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            doc.add_heading(heading_match.group(2), level=level)
            i += 1
            continue

        if stripped.strip().startswith("|"):
            rows, i = parse_table(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for r, row in enumerate(rows):
                    for c, cell_text in enumerate(row):
                        cell = table.rows[r].cells[c]
                        cell.paragraphs[0].text = ""
                        add_inline_runs(cell.paragraphs[0], cell_text)
                        if r == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
                doc.add_paragraph()
            continue

        bullet_match = re.match(r"^(\s*)-\s+(.*)$", stripped)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(2)
            # Continuation lines of a wrapped bullet are indented further
            # without their own leading "-"; fold them into one paragraph.
            j = i + 1
            while (
                j < len(lines)
                and lines[j].strip()
                and not re.match(r"^(\s*)-\s+", lines[j])
                and not lines[j].strip().startswith("|")
                and not re.match(r"^#{1,3}\s", lines[j])
            ):
                text += " " + lines[j].strip()
                j += 1
            p = doc.add_paragraph(style="List Bullet" if indent == 0 else "List Bullet 2")
            add_inline_runs(p, text)
            i = j
            continue

        # Plain paragraph: fold wrapped lines until a blank line or new block.
        text = stripped
        j = i + 1
        while (
            j < len(lines)
            and lines[j].strip()
            and not re.match(r"^(\s*)-\s+", lines[j])
            and not lines[j].strip().startswith("|")
            and not re.match(r"^#{1,3}\s", lines[j])
        ):
            text += " " + lines[j].strip()
            j += 1
        p = doc.add_paragraph()
        add_inline_runs(p, text)
        i = j


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("usage: python scripts/md_to_docx.py <src.md> <dst.docx>")
    src, dst = Path(sys.argv[1]), Path(sys.argv[2])

    text = src.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    build(doc, lines)
    doc.save(dst)
    print(f"wrote {dst}")


if __name__ == "__main__":
    main()
